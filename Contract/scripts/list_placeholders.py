from __future__ import annotations

import re
from pathlib import Path

import docx


def main() -> None:
    template = Path(r"F:\Dev\Contract\templates\HDQTGAN_PN_MR_template.docx")
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")

    d = docx.Document(str(template))
    found: set[str] = set()

    for par in d.paragraphs:
        found.update(re.findall(r"\{\{[^}]+\}\}", par.text))

    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    found.update(re.findall(r"\{\{[^}]+\}\}", par.text))

    print(f"template: {template}")
    print(f"placeholders: {len(found)}")
    for x in sorted(found):
        print(x)


if __name__ == "__main__":
    main()
